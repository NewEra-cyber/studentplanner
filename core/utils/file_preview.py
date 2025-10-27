# core/utils/file_preview.py
import os
import mimetypes
from django.conf import settings
import docx
import PyPDF2
from PIL import Image
import pandas as pd
import markdown
from bs4 import BeautifulSoup
import chardet
import magic

class FilePreviewGenerator:
    @staticmethod
    def generate_preview(file_path, max_length=10000):
        """Enhanced preview for various file types with better formatting"""
        if not os.path.exists(file_path):
            return "File not found"
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            # Text files
            if file_extension in ['.txt', '.md', '.py', '.js', '.css', '.html', '.json', '.xml']:
                return FilePreviewGenerator._preview_text_file(file_path, max_length)
            
            # PDF files
            elif file_extension == '.pdf':
                return FilePreviewGenerator._preview_pdf(file_path, max_length)
            
            # Word documents
            elif file_extension in ['.docx', '.doc']:
                return FilePreviewGenerator._preview_docx(file_path, max_length)
            
            # Excel files
            elif file_extension in ['.xlsx', '.xls', '.csv']:
                return FilePreviewGenerator._preview_excel(file_path, max_length)
            
            # Images
            elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg']:
                return FilePreviewGenerator._preview_image(file_path)
            
            # PowerPoint
            elif file_extension in ['.pptx', '.ppt']:
                return FilePreviewGenerator._preview_powerpoint(file_path)
            
            else:
                return f"📄 File Type: {file_extension.upper()}\n\nPreview available for text-based files. Download to view full content."
                
        except Exception as e:
            return f"❌ Error generating preview: {str(e)}"
    
    @staticmethod
    def _preview_text_file(file_path, max_length):
        """Enhanced text file preview with syntax highlighting support"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            
            # Read content
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read(max_length)
                
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # Format based on file type
            if file_extension == '.md':
                # Convert markdown to HTML for better display
                html_content = markdown.markdown(content[:2000])
                return f"📝 **Markdown Preview**\n\n{html_content}"
            
            elif file_extension == '.html':
                # Extract text from HTML
                soup = BeautifulSoup(content, 'html.parser')
                text = soup.get_text()
                return f"🌐 **HTML Content**\n\n{text[:max_length]}"
            
            elif file_extension in ['.py', '.js', '.css']:
                # Code files with line numbers
                lines = content.split('\n')
                numbered_lines = [f"{i+1:3d} | {line}" for i, line in enumerate(lines[:50])]
                return f"💻 **Code Preview** ({file_extension})\n\n" + '\n'.join(numbered_lines)
            
            else:
                # Regular text files
                if len(content) == max_length:
                    content += "\n\n... (content truncated)"
                return f"📄 **Text Content**\n\n{content}"
                
        except Exception as e:
            return f"Error reading text file: {str(e)}"
    
    @staticmethod
    def _preview_pdf(file_path, max_length):
        """Enhanced PDF preview with better text extraction"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                
                # Extract from first 5 pages or until max_length
                for i, page in enumerate(pdf_reader.pages[:5]):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"--- Page {i+1} ---\n{page_text}\n\n"
                    if len(text) > max_length:
                        text = text[:max_length] + "\n\n... (content truncated)"
                        break
                
                if text.strip():
                    return f"📕 **PDF Preview** ({len(pdf_reader.pages)} pages)\n\n{text}"
                else:
                    return "📕 **PDF Document**\n\nThis appears to be a scanned PDF or image-based document. No extractable text found."
                    
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    @staticmethod
    def _preview_docx(file_path, max_length):
        """Enhanced DOCX preview with structured content"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs[:100]):  # First 100 paragraphs
                if paragraph.text.strip():
                    # Detect heading styles
                    if paragraph.style.name.startswith('Heading'):
                        full_text.append(f"\n# {paragraph.text}\n")
                    else:
                        full_text.append(paragraph.text)
                    
                    if len('\n'.join(full_text)) > max_length:
                        full_text.append("\n\n... (content truncated)")
                        break
            
            # Extract tables
            for table in doc.tables[:3]:  # First 3 tables
                full_text.append("\n--- Table ---")
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_data))
                full_text.append("--- End Table ---\n")
            
            content = '\n'.join(full_text)
            return f"📘 **Word Document Preview**\n\n{content}"
            
        except Exception as e:
            return f"Error reading Word document: {str(e)}"
    
    @staticmethod
    def _preview_excel(file_path, max_length):
        """Enhanced Excel/CSV preview with data summary"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.csv':
                df = pd.read_csv(file_path, nrows=15)
                file_type = "CSV"
            else:
                df = pd.read_excel(file_path, nrows=15)
                file_type = "Excel"
            
            preview = f"📊 **{file_type} File Preview**\n\n"
            preview += f"Data Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
            preview += "First 10 rows:\n\n"
            
            # Format the dataframe for better display
            preview += df.head(10).to_string(max_rows=10, max_cols=8, index=True)
            
            # Add column info
            preview += f"\n\nColumns: {', '.join(df.columns.astype(str))}"
            
            return preview
            
        except Exception as e:
            return f"Error reading spreadsheet: {str(e)}"
    
    @staticmethod
    def _preview_image(file_path):
        """Enhanced image preview with metadata"""
        try:
            with Image.open(file_path) as img:
                file_size = os.path.getsize(file_path)
                size_mb = file_size / (1024 * 1024)
                
                preview = f"🖼️ **Image Preview**\n\n"
                preview += f"Format: {img.format or 'Unknown'}\n"
                preview += f"Dimensions: {img.size[0]} × {img.size[1]} pixels\n"
                preview += f"Color Mode: {img.mode}\n"
                preview += f"File Size: {size_mb:.2f} MB\n\n"
                preview += "Image content cannot be displayed in text preview. Download to view."
                
                return preview
        except Exception as e:
            return f"Error reading image: {str(e)}"
    
    @staticmethod
    def _preview_powerpoint(file_path):
        """Basic PowerPoint preview"""
        return "📊 **PowerPoint Presentation**\n\nSlide deck content preview not available in text format. Download to view the presentation."
    
    @staticmethod
    def get_file_type_info(filename):
        """Get detailed file type information"""
        file_extension = os.path.splitext(filename)[1].lower()
        
        type_info = {
            'icon': 'fas fa-file',
            'color': 'default',
            'type_name': 'Document',
            'preview_supported': True
        }
        
        file_types = {
            '.txt': {'icon': 'fas fa-file-alt', 'color': 'txt', 'type_name': 'Text File'},
            '.pdf': {'icon': 'fas fa-file-pdf', 'color': 'pdf', 'type_name': 'PDF Document'},
            '.doc': {'icon': 'fas fa-file-word', 'color': 'doc', 'type_name': 'Word Document'},
            '.docx': {'icon': 'fas fa-file-word', 'color': 'doc', 'type_name': 'Word Document'},
            '.xls': {'icon': 'fas fa-file-excel', 'color': 'excel', 'type_name': 'Excel Spreadsheet'},
            '.xlsx': {'icon': 'fas fa-file-excel', 'color': 'excel', 'type_name': 'Excel Spreadsheet'},
            '.csv': {'icon': 'fas fa-file-csv', 'color': 'csv', 'type_name': 'CSV File'},
            '.ppt': {'icon': 'fas fa-file-powerpoint', 'color': 'powerpoint', 'type_name': 'PowerPoint'},
            '.pptx': {'icon': 'fas fa-file-powerpoint', 'color': 'powerpoint', 'type_name': 'PowerPoint'},
            '.html': {'icon': 'fas fa-file-code', 'color': 'html', 'type_name': 'HTML File'},
            '.htm': {'icon': 'fas fa-file-code', 'color': 'html', 'type_name': 'HTML File'},
            '.py': {'icon': 'fas fa-file-code', 'color': 'code', 'type_name': 'Python Script'},
            '.js': {'icon': 'fas fa-file-code', 'color': 'code', 'type_name': 'JavaScript'},
            '.css': {'icon': 'fas fa-file-code', 'color': 'code', 'type_name': 'Stylesheet'},
            '.jpg': {'icon': 'fas fa-file-image', 'color': 'image', 'type_name': 'JPEG Image'},
            '.jpeg': {'icon': 'fas fa-file-image', 'color': 'image', 'type_name': 'JPEG Image'},
            '.png': {'icon': 'fas fa-file-image', 'color': 'image', 'type_name': 'PNG Image'},
            '.gif': {'icon': 'fas fa-file-image', 'color': 'image', 'type_name': 'GIF Image'},
            '.zip': {'icon': 'fas fa-file-archive', 'color': 'archive', 'type_name': 'ZIP Archive'},
            '.rar': {'icon': 'fas fa-file-archive', 'color': 'archive', 'type_name': 'RAR Archive'},
        }
        
        return type_info.update(file_types.get(file_extension, {})) or type_info